#!/usr/bin/env python3
"""
Content Reviewer — World-class editorial pass with structured notes.

Analyzes content pieces for polish and refinement. Does NOT edit text. Produces:
  1. A readable editorial document (docx)
  2. A spreadsheet with every note organized for action

Usage:
    export ANTHROPIC_API_KEY=sk-ant-...
    python3 reviewer.py
    python3 reviewer.py --items 3,12,20
    python3 reviewer.py --output my_review

Legacy tool. Use agent_loop.py with review mode instead.
Framework: Lyra Labs, 2026
"""

import json
import os
import sys
import re
import time
import argparse
from pathlib import Path
from datetime import datetime, timezone

sys.path.insert(0, str(Path(__file__).parent))

SCRIPT_DIR = Path(__file__).parent

# ═══════════════════════════════════════════════════════════
# REVIEWER SYSTEM — World-class editorial analysis
# ═══════════════════════════════════════════════════════════

REVIEWER_SYSTEM = """You are a world-class editor doing a final polish pass on content.

Your goal is POLISH, not restructuring.

EDITORIAL PHILOSOPHY:
- You are here to refine, not fundamentally restructure
- Every note should take under 5 minutes to implement
- If something works well, say nothing. Silence is approval.
- You only flag things that would make a discerning reader pause
- You respect the author's voice and intent
- Content structure is intentional. Don't question it.

WHAT YOU LOOK FOR:

1. SENTENCE-LEVEL POLISH
   - Sentences that could lose 3-5 words without losing meaning
   - Repeated words within the same paragraph (not across chapters — that's intentional)
   - Awkward rhythm — sentences that stumble when read aloud
   - A strong sentence buried in a weak paragraph

2. FIELD STORIES (the lived examples)
   - Does the story have a NAME or PLACE? (If not, it reads as generic)
   - Can you SEE it? (One sensory detail makes the difference)
   - Does it earn its length? (A 200-word story that could land in 80 words)

3. PARAGRAPH TRANSITIONS
   - Any place where two paragraphs feel stapled together rather than flowing
   - A paragraph that starts strong but trails off
   - A place where one paragraph should become two (or two should become one)

4. THE OPENING AND CLOSING
   - Does the chapter's first sentence pull you in?
   - Does the last sentence land with weight? (Best closers feel like a mic drop or a quiet exhale)
   - Would a reader turn the page?

5. QUOTABILITY
   - Flag any line that's ALMOST a standalone quote but needs 2-3 words trimmed to get there
   - Note lines that ARE already perfect quotes (these are assets)

6. MINOR CLEANUPS
   - Tense inconsistency within a passage
   - Pronoun ambiguity (does "it" clearly refer to something?)
   - Any sentence where you had to read twice to understand

WHAT YOU DO NOT DO:
- Do NOT question whether a chapter should exist (that decision is made)
- Do NOT suggest restructuring the chapter order
- Do NOT flag the philosophical/spiritual tone as a problem — it IS the book
- Do NOT suggest adding more content. Less is more at this stage.
- Do NOT rewrite passages. Just note what's off and trust the author to fix it.
- Do NOT flag things that are working. No "this is great" notes — silence IS praise.

OUTPUT FORMAT — respond with ONLY valid JSON:
{
  "chapter_impression": "One sentence — your gut reaction as a reader. What lands, what lingers.",
  "opening_verdict": "HOOKS / SLOW / FLAT — then one sentence why",
  "closing_verdict": "LANDS / TRAILS / FLAT — then one sentence why",
  "best_line": "The single best sentence in this chapter, verbatim",
  "notes": [
    {
      "location": "Quote 5-10 words from the text where the note applies",
      "note": "What you'd say if you were sitting next to the author with a pencil. Keep it conversational, specific, actionable.",
      "type": "trim|rhythm|transition|story|opening|closing|clarity|quote_polish|cleanup",
      "effort": "30sec|2min|5min"
    }
  ],
  "word_trim_estimate": "How many words could this chapter probably lose? Give a number and one-line rationale.",
  "page_turner_score": "1-10: Would a reader keep going? 8+ means yes. Be honest but fair."
}

IMPORTANT: Keep notes to genuinely useful observations. A chapter with 3 sharp notes is better than one with 15 generic ones. If a chapter is clean, it's OK to return just 1-2 notes. An empty notes array means you think it's ready for print — use that sparingly but use it."""

# ═══════════════════════════════════════════════════════════
# CROSS-CHAPTER FINAL NOTES
# ═══════════════════════════════════════════════════════════

CROSS_CHAPTER_SYSTEM = """You are doing a final cross-chapter read of the manuscript. You've just read every chapter. Now step back and note patterns across the WHOLE book.

You are looking for:
1. REPEATED PHRASES that appear in 3+ chapters (exact or near-exact wording, not thematic repetition — themes are intentional)
2. Any chapter that feels like it covers the SAME GROUND as another (not theme, but actual overlapping explanations)
3. Pacing across the book — does any movement feel long? Does the energy dip anywhere?
4. The book's BEST 10 LINES — the ones that would go on the back cover or in a keynote

Respond with ONLY valid JSON:
{
  "repeated_phrases": [
    {"phrase": "the exact repeated wording", "chapters": ["Ch 3", "Ch 12", "Ch 20"], "suggestion": "keep in Ch X, trim from others"}
  ],
  "overlapping_chapters": [
    {"chapters": ["Ch X", "Ch Y"], "overlap": "What they both explain", "suggestion": "one sentence fix"}
  ],
  "pacing_notes": [
    {"location": "Movement X / Ch Y-Z", "note": "What's happening with pacing here"}
  ],
  "best_10_lines": [
    {"line": "exact quote", "chapter": "Ch X"}
  ],
  "overall_word_trim": "Total words the book could probably lose. Number + rationale."
}"""


# ═══════════════════════════════════════════════════════════
# API CALL
# ═══════════════════════════════════════════════════════════

def call_api(prompt, system, api_key, model, max_tokens=4096):
    """Call Anthropic API."""
    import urllib.request
    url = "https://api.anthropic.com/v1/messages"
    headers = {
        "Content-Type": "application/json",
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
    }
    body = json.dumps({
        "model": model,
        "max_tokens": max_tokens,
        "system": system,
        "messages": [{"role": "user", "content": prompt}],
    }).encode()

    req = urllib.request.Request(url, data=body, headers=headers, method="POST")
    for attempt in range(3):
        try:
            with urllib.request.urlopen(req, timeout=120) as resp:
                data = json.loads(resp.read().decode())
                return data["content"][0]["text"]
        except Exception as e:
            if attempt < 2:
                print(f"    API retry ({attempt+1}/2): {str(e)[:80]}")
                time.sleep(5 * (attempt + 1))
            else:
                raise


def review_chapter(title, text, api_key, model):
    """Send one chapter through the reviewer's lens."""
    word_count = len(text.split())
    prompt = f"""CHAPTER: {title}
WORD COUNT: {word_count}

---

{text}

---

Read this chapter as a world-class editor on the final pass before print. Return your editorial notes as JSON."""

    for attempt in range(2):
        try:
            raw = call_api(prompt, REVIEWER_SYSTEM, api_key, model, max_tokens=4096)
            # Extract JSON from response
            raw = raw.strip()
            if raw.startswith("```"):
                raw = re.sub(r'^```(?:json)?\n?', '', raw)
                raw = re.sub(r'\n?```$', '', raw)
            return json.loads(raw)
        except json.JSONDecodeError:
            if attempt == 0:
                print(f"    JSON parse retry...")
                time.sleep(3)
            else:
                print(f"    WARNING: Could not parse JSON for {title}")
                return {"error": "JSON parse failed", "raw": raw[:500]}
        except Exception as e:
            print(f"    ERROR: {str(e)[:100]}")
            return {"error": str(e)[:200]}


def cross_chapter_review(reviews, api_key, model):
    """Run the cross-chapter analysis after all individual reviews are done."""
    # Build a summary of each chapter for the cross-chapter reader
    summaries = []
    for title, review in reviews.items():
        if "error" in review:
            continue
        best = review.get("best_line", "")
        impression = review.get("chapter_impression", "")
        notes_count = len(review.get("notes", []))
        trim = review.get("word_trim_estimate", "")
        summaries.append(f"**{title}** — {impression}\n  Best line: \"{best}\"\n  Notes: {notes_count} | Trim estimate: {trim}")

    prompt = f"""Here are editorial summaries of every chapter in the manuscript:

{chr(10).join(summaries)}

Now do your cross-chapter analysis. Return JSON."""

    try:
        raw = call_api(prompt, CROSS_CHAPTER_SYSTEM, api_key, model, max_tokens=4096)
        raw = raw.strip()
        if raw.startswith("```"):
            raw = re.sub(r'^```(?:json)?\n?', '', raw)
            raw = re.sub(r'\n?```$', '', raw)
        return json.loads(raw)
    except Exception as e:
        print(f"  Cross-chapter analysis error: {str(e)[:100]}")
        return {"error": str(e)[:200]}


# ═══════════════════════════════════════════════════════════
# BUILD THE DOCX — The readable editorial letter
# ═══════════════════════════════════════════════════════════

def build_editorial_letter(reviews, cross_review, output_path):
    """Build a readable Word doc — editorial letter to the author."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  Installing python-docx...")
        import subprocess
        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "--break-system-packages", "-q"])
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("Manuscript — Final Editorial Notes", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    date_para = doc.add_paragraph(f"Prepared: {datetime.now().strftime('%B %d, %Y')}")
    date_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0xAA)

    doc.add_paragraph(
        "These are final-pass notes — the kind a world-class editor leaves in the margins "
        "before a book goes to print. Every note should take under 5 minutes to implement. "
        "Silence on a chapter means it's ready."
    )
    doc.add_paragraph("")

    # Stats
    total_notes = 0
    chapters_clean = 0
    chapters_reviewed = 0

    for title_text, review in reviews.items():
        if "error" in review:
            continue
        chapters_reviewed += 1
        notes = review.get("notes", [])
        total_notes += len(notes)
        if len(notes) <= 1:
            chapters_clean += 1

    stats = doc.add_paragraph()
    stats.add_run(f"{chapters_reviewed} chapters reviewed").bold = True
    stats.add_run(f" · {total_notes} notes total · {chapters_clean} chapters near-clean")

    doc.add_paragraph("")

    # ── Per-chapter sections ──
    for ch_title, review in reviews.items():
        if "error" in review:
            continue

        notes = review.get("notes", [])
        impression = review.get("chapter_impression", "")
        opening = review.get("opening_verdict", "")
        closing = review.get("closing_verdict", "")
        best = review.get("best_line", "")
        trim = review.get("word_trim_estimate", "")
        score = review.get("page_turner_score", "")

        # Chapter heading
        h = doc.add_heading(ch_title, level=2)

        # Impression + score
        if impression:
            p = doc.add_paragraph()
            p.add_run(impression).italic = True
            if score:
                p.add_run(f"  [Page-turner: {score}/10]")

        # Opening / Closing
        if opening or closing:
            oc = doc.add_paragraph()
            if opening:
                oc.add_run(f"Opening: {opening}").font.size = Pt(10)
            if opening and closing:
                oc.add_run("  ·  ")
            if closing:
                oc.add_run(f"Closing: {closing}").font.size = Pt(10)

        # Best line
        if best:
            bp = doc.add_paragraph()
            bp.add_run("Best line: ").bold = True
            bp.add_run(f'"{best}"').italic = True

        # Trim estimate
        if trim:
            tp = doc.add_paragraph()
            tp.add_run(f"Trim estimate: {trim}").font.color.rgb = RGBColor(0x55, 0x55, 0x66)

        # Notes
        if not notes:
            doc.add_paragraph("No notes — ready for print.").runs[0].font.color.rgb = RGBColor(0x22, 0x88, 0x22)
        else:
            for i, note in enumerate(notes):
                loc = note.get("location", "")
                text = note.get("note", "")
                ntype = note.get("type", "")
                effort = note.get("effort", "")

                np = doc.add_paragraph()
                np.add_run(f"{i+1}. ").bold = True
                if loc:
                    run = np.add_run(f'"{loc}" ')
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x8B, 0x7E, 0xC8)
                np.add_run(f"— {text}")
                if ntype or effort:
                    tag = np.add_run(f"  [{ntype}" + (f", {effort}" if effort else "") + "]")
                    tag.font.size = Pt(9)
                    tag.font.color.rgb = RGBColor(0x88, 0x88, 0xAA)

        doc.add_paragraph("")  # spacing

    # ── Cross-chapter section ──
    if cross_review and "error" not in cross_review:
        doc.add_heading("Cross-Chapter Notes", level=1)

        # Repeated phrases
        reps = cross_review.get("repeated_phrases", [])
        if reps:
            doc.add_heading("Repeated Phrases", level=3)
            for r in reps:
                p = doc.add_paragraph()
                p.add_run(f'"{r.get("phrase", "")}"').italic = True
                p.add_run(f' — appears in {", ".join(r.get("chapters", []))}')
                sugg = r.get("suggestion", "")
                if sugg:
                    p.add_run(f"\n  → {sugg}").font.color.rgb = RGBColor(0x55, 0x55, 0x66)

        # Overlaps
        overlaps = cross_review.get("overlapping_chapters", [])
        if overlaps:
            doc.add_heading("Overlapping Chapters", level=3)
            for o in overlaps:
                p = doc.add_paragraph()
                p.add_run(f'{", ".join(o.get("chapters", []))}').bold = True
                p.add_run(f' — {o.get("overlap", "")}')
                sugg = o.get("suggestion", "")
                if sugg:
                    p.add_run(f"\n  → {sugg}").font.color.rgb = RGBColor(0x55, 0x55, 0x66)

        # Pacing
        pacing = cross_review.get("pacing_notes", [])
        if pacing:
            doc.add_heading("Pacing Notes", level=3)
            for pn in pacing:
                p = doc.add_paragraph()
                p.add_run(f'{pn.get("location", "")}').bold = True
                p.add_run(f' — {pn.get("note", "")}')

        # Best 10 lines
        best_lines = cross_review.get("best_10_lines", [])
        if best_lines:
            doc.add_heading("The Book's 10 Best Lines", level=2)
            for bl in best_lines:
                p = doc.add_paragraph()
                p.add_run(f'"{bl.get("line", "")}"').italic = True
                ch = bl.get("chapter", "")
                if ch:
                    p.add_run(f"  — {ch}").font.color.rgb = RGBColor(0x88, 0x88, 0xAA)

        # Overall trim
        trim = cross_review.get("overall_word_trim", "")
        if trim:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            p.add_run(f"Overall trim estimate: {trim}").bold = True

    doc.save(str(output_path))
    print(f"  Editorial letter saved: {output_path}")


# ═══════════════════════════════════════════════════════════
# BUILD THE SPREADSHEET
# ═══════════════════════════════════════════════════════════

def build_spreadsheet(reviews, cross_review, chapters_data, output_path):
    """Build the action spreadsheet."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        import subprocess
        subprocess.run([sys.executable, "-m", "pip", "install", "openpyxl", "--break-system-packages", "-q"])
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    hdr_font = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
    hdr_fill = PatternFill(start_color="8B7EC8", end_color="8B7EC8", fill_type="solid")
    wrap = Alignment(wrap_text=True, vertical="top")
    thin_border = Border(
        left=Side(style="thin", color="E0DDF0"),
        right=Side(style="thin", color="E0DDF0"),
        top=Side(style="thin", color="E0DDF0"),
        bottom=Side(style="thin", color="E0DDF0"),
    )

    def style_header(ws, row=1):
        for cell in ws[row]:
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = wrap
            cell.border = thin_border

    # ── Sheet 1: Chapter Overview ──
    ws1 = wb.active
    ws1.title = "Chapter Overview"
    ws1.append(["Chapter", "Words", "Page-Turner", "Opening", "Closing", "Notes", "Trim Est.", "Impression", "Best Line"])
    style_header(ws1)

    green_fill = PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid")
    yellow_fill = PatternFill(start_color="FFF8E1", end_color="FFF8E1", fill_type="solid")

    for ch_title, review in reviews.items():
        if "error" in review:
            ws1.append([ch_title, "", "", "", "", "ERROR", "", review.get("error", "")])
            continue

        wc = chapters_data.get(ch_title, {}).get("words", 0)
        score = review.get("page_turner_score", "")
        opening = review.get("opening_verdict", "")
        closing = review.get("closing_verdict", "")
        notes = len(review.get("notes", []))
        trim = review.get("word_trim_estimate", "")
        impression = review.get("chapter_impression", "")
        best = review.get("best_line", "")

        ws1.append([ch_title, wc, score, opening, closing, notes, trim, impression, best])

        # Color code
        row = ws1.max_row
        if notes <= 1:
            for cell in ws1[row]:
                cell.fill = green_fill
        elif notes >= 5:
            for cell in ws1[row]:
                cell.fill = yellow_fill

    for cell in ws1[ws1.max_row + 1 if ws1.max_row else 1]:
        cell.border = thin_border

    ws1.column_dimensions["A"].width = 40
    ws1.column_dimensions["B"].width = 8
    ws1.column_dimensions["C"].width = 10
    ws1.column_dimensions["D"].width = 25
    ws1.column_dimensions["E"].width = 25
    ws1.column_dimensions["F"].width = 8
    ws1.column_dimensions["G"].width = 20
    ws1.column_dimensions["H"].width = 50
    ws1.column_dimensions["I"].width = 50

    # ── Sheet 2: All Notes ──
    ws2 = wb.create_sheet("All Notes")
    ws2.append(["Chapter", "Location", "Note", "Type", "Effort"])
    style_header(ws2)

    for ch_title, review in reviews.items():
        if "error" in review:
            continue
        for note in review.get("notes", []):
            ws2.append([
                ch_title,
                note.get("location", ""),
                note.get("note", ""),
                note.get("type", ""),
                note.get("effort", ""),
            ])

    ws2.column_dimensions["A"].width = 35
    ws2.column_dimensions["B"].width = 30
    ws2.column_dimensions["C"].width = 60
    ws2.column_dimensions["D"].width = 15
    ws2.column_dimensions["E"].width = 10

    # Apply wrapping to all data cells
    for row in ws2.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = wrap
            cell.border = thin_border

    # ── Sheet 3: Quick Wins (30sec and 2min notes) ──
    ws3 = wb.create_sheet("Quick Wins")
    ws3.append(["Chapter", "Location", "Note", "Type", "Effort"])
    style_header(ws3)

    for ch_title, review in reviews.items():
        if "error" in review:
            continue
        for note in review.get("notes", []):
            if note.get("effort", "") in ("30sec", "2min"):
                ws3.append([
                    ch_title,
                    note.get("location", ""),
                    note.get("note", ""),
                    note.get("type", ""),
                    note.get("effort", ""),
                ])

    ws3.column_dimensions["A"].width = 35
    ws3.column_dimensions["B"].width = 30
    ws3.column_dimensions["C"].width = 60
    ws3.column_dimensions["D"].width = 15
    ws3.column_dimensions["E"].width = 10

    # ── Sheet 4: Best Lines ──
    ws4 = wb.create_sheet("Best Lines")
    ws4.append(["Chapter", "Line"])
    style_header(ws4)

    for ch_title, review in reviews.items():
        if "error" in review:
            continue
        best = review.get("best_line", "")
        if best:
            ws4.append([ch_title, best])

    # Add cross-chapter best 10
    if cross_review and "error" not in cross_review:
        ws4.append([])
        ws4.append(["— TOP 10 ACROSS THE BOOK —", ""])
        for bl in cross_review.get("best_10_lines", []):
            ws4.append([bl.get("chapter", ""), bl.get("line", "")])

    ws4.column_dimensions["A"].width = 35
    ws4.column_dimensions["B"].width = 80

    # ── Sheet 5: Cross-Chapter ──
    if cross_review and "error" not in cross_review:
        ws5 = wb.create_sheet("Cross-Chapter")

        ws5.append(["REPEATED PHRASES"])
        ws5.append(["Phrase", "Chapters", "Suggestion"])
        style_header(ws5, ws5.max_row)
        for r in cross_review.get("repeated_phrases", []):
            ws5.append([r.get("phrase", ""), ", ".join(r.get("chapters", [])), r.get("suggestion", "")])

        ws5.append([])
        ws5.append(["OVERLAPPING CHAPTERS"])
        row_start = ws5.max_row + 1
        ws5.append(["Chapters", "Overlap", "Suggestion"])
        style_header(ws5, ws5.max_row)
        for o in cross_review.get("overlapping_chapters", []):
            ws5.append([", ".join(o.get("chapters", [])), o.get("overlap", ""), o.get("suggestion", "")])

        ws5.append([])
        ws5.append(["PACING NOTES"])
        ws5.append(["Location", "Note"])
        style_header(ws5, ws5.max_row)
        for pn in cross_review.get("pacing_notes", []):
            ws5.append([pn.get("location", ""), pn.get("note", "")])

        trim = cross_review.get("overall_word_trim", "")
        if trim:
            ws5.append([])
            ws5.append([f"Overall trim estimate: {trim}"])

        ws5.column_dimensions["A"].width = 35
        ws5.column_dimensions["B"].width = 50
        ws5.column_dimensions["C"].width = 50

    wb.save(str(output_path))
    print(f"  Spreadsheet saved: {output_path}")


# ═══════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="Editorial Reviewer — Final editorial pass")
    parser.add_argument("--chapters-json", default=str(SCRIPT_DIR / "chapters_all.json"),
                        help="Path to chapters_all.json")
    parser.add_argument("--chapters", default="",
                        help="Comma-separated chapter numbers to review (default: all)")
    parser.add_argument("--anthropic-key", default=os.environ.get("ANTHROPIC_API_KEY", ""))
    parser.add_argument("--model", default="claude-sonnet-4-20250514")
    parser.add_argument("--output", default="Editorial_Notes")
    args = parser.parse_args()

    api_key = args.anthropic_key
    if not api_key:
        print("ERROR: Set ANTHROPIC_API_KEY")
        sys.exit(1)

    # Load chapters
    chapters_path = Path(args.chapters_json)
    if not chapters_path.exists():
        print(f"ERROR: {chapters_path} not found")
        sys.exit(1)

    with open(chapters_path) as f:
        chapters_data = json.load(f)

    print(f"\n{'='*70}")
    print(f"  EDITORIAL REVIEWER — Final Editorial Pass")
    print(f"  Manuscript Review")
    print(f"  {len(chapters_data)} sections loaded")
    print(f"{'='*70}\n")

    # Filter to actual chapters (skip dedications, style guides, etc.)
    chapter_filter = None
    if args.chapters:
        chapter_filter = [c.strip() for c in args.chapters.split(",")]

    review_chapters = {}
    for title, data in chapters_data.items():
        text = data.get("text", "") if isinstance(data, dict) else str(data)
        words = len(text.split()) if text else 0

        # Skip very short sections and non-chapter content
        if words < 200:
            continue

        # If specific chapters requested, filter
        if chapter_filter:
            ch_num = re.search(r'Chapter\s+(\d+)', title)
            if ch_num and ch_num.group(1) in chapter_filter:
                review_chapters[title] = text
            elif any(f.lower() in title.lower() for f in chapter_filter):
                review_chapters[title] = text
        else:
            review_chapters[title] = text

    print(f"  Reviewing {len(review_chapters)} chapters...\n")

    # ── Review each chapter (with resume support) ──
    progress_path = output_dir / f"{output_base}_progress.json"
    reviews = {}

    # Load existing progress if available
    if progress_path.exists():
        try:
            with open(progress_path) as f:
                reviews = json.load(f)
            print(f"  Resuming from {len(reviews)} previously reviewed chapters...\n")
        except:
            reviews = {}

    for i, (title, text) in enumerate(review_chapters.items()):
        # Skip already reviewed chapters
        if title in reviews and "error" not in reviews[title]:
            print(f"  [{i+1}/{len(review_chapters)}] {title} — already reviewed, skipping")
            continue

        word_count = len(text.split())
        print(f"  [{i+1}/{len(review_chapters)}] {title} ({word_count} words)")

        review = review_chapter(title, text, api_key, args.model)
        reviews[title] = review

        # Print summary
        if "error" not in review:
            notes = review.get("notes", [])
            score = review.get("page_turner_score", "")
            impression = review.get("chapter_impression", "")
            print(f"    → {len(notes)} notes | Page-turner: {score}/10")
            if impression:
                print(f"    → {impression[:80]}")
        else:
            print(f"    → ERROR: {review.get('error', '')[:80]}")

        # Save progress after each chapter
        with open(progress_path, "w") as f:
            json.dump(reviews, f, indent=2)

        print()
        time.sleep(1)  # Rate limiting

    # ── Cross-chapter analysis ──
    print(f"\n  Running cross-chapter analysis...")
    cross_review = cross_chapter_review(reviews, api_key, args.model)

    if "error" not in cross_review:
        reps = cross_review.get("repeated_phrases", [])
        overlaps = cross_review.get("overlapping_chapters", [])
        best = cross_review.get("best_10_lines", [])
        print(f"    → {len(reps)} repeated phrases, {len(overlaps)} overlaps, {len(best)} best lines")
    else:
        print(f"    → ERROR: {cross_review.get('error', '')[:80]}")

    # ── Save outputs ──
    output_dir = SCRIPT_DIR
    output_base = args.output

    # JSON backup
    json_path = output_dir / f"{output_base}.json"
    with open(json_path, "w") as f:
        json.dump({"reviews": reviews, "cross_chapter": cross_review}, f, indent=2)
    print(f"\n  JSON backup: {json_path}")

    # Docx editorial letter
    docx_path = output_dir / f"{output_base}.docx"
    build_editorial_letter(reviews, cross_review, docx_path)

    # Spreadsheet
    xlsx_path = output_dir / f"{output_base}.xlsx"
    build_spreadsheet(reviews, cross_review, chapters_data, xlsx_path)

    # ── Summary ──
    total_notes = sum(len(r.get("notes", [])) for r in reviews.values() if "error" not in r)
    quick_wins = sum(
        1 for r in reviews.values() if "error" not in r
        for n in r.get("notes", []) if n.get("effort", "") in ("30sec", "2min")
    )

    print(f"\n{'='*70}")
    print(f"  DONE")
    print(f"  {len(reviews)} chapters reviewed")
    print(f"  {total_notes} total notes ({quick_wins} quick wins)")
    print(f"  Editorial letter: {docx_path}")
    print(f"  Spreadsheet: {xlsx_path}")
    print(f"{'='*70}\n")


if __name__ == "__main__":
    main()
